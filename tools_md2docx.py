"""Converteste reports/lucrare.md in reports/lucrare.docx.

Formulele LaTeX sunt randate ca imagini cu matplotlib (mathtext), figurile
referite sunt inserate efectiv, iar tabelele markdown devin tabele Word.
"""

import hashlib
import os
import re
import sys

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

MD = "reports/lucrare.md"
OUT = "reports/lucrare.docx"
FORMULE_DIR = "reports/figures/formule"

GRI = RGBColor(0x44, 0x44, 0x44)


# --------------------------------------------------------------- formule
def randeaza_formula(latex, display=True):
    """Randeaza o formula LaTeX ca PNG si returneaza calea."""
    os.makedirs(FORMULE_DIR, exist_ok=True)
    cheie = hashlib.md5((latex + str(display)).encode("utf-8")).hexdigest()[:12]
    cale = os.path.join(FORMULE_DIR, f"f_{cheie}.png")
    if os.path.exists(cale):
        return cale

    marime = 15 if display else 11
    fig = plt.figure(figsize=(0.01, 0.01))
    fig.text(0, 0, f"${latex}$", fontsize=marime)
    try:
        fig.savefig(cale, dpi=220, bbox_inches="tight", pad_inches=0.06,
                    transparent=False, facecolor="white")
    except Exception as e:
        plt.close(fig)
        print(f"    ! formula neredabila: {latex[:50]} ({e})")
        return None
    plt.close(fig)
    return cale


# --------------------------------------------------------- text inline
def scrie_inline(paragraf, text, baza_bold=False, italic=False):
    """Scrie text cu **bold**, *italic*, `cod` si $formule$ inline.

    Se apeleaza recursiv pentru continutul dintre ** **, altfel un `cod`
    aflat in interiorul unui text ingrosat ar ramane cu backtick-uri.
    """
    bucati = re.split(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*|`[^`]+?`|\$[^$]+?\$)", text)
    for b in bucati:
        if not b:
            continue
        if b.startswith("**") and b.endswith("**"):
            scrie_inline(paragraf, b[2:-2], baza_bold=True, italic=italic)
            continue
        elif b.startswith("`") and b.endswith("`"):
            r = paragraf.add_run(b[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        elif b.startswith("$") and b.endswith("$") and len(b) > 2:
            cale = randeaza_formula(b[1:-1], display=False)
            if cale:
                paragraf.add_run().add_picture(cale, height=Pt(11))
            else:
                r = paragraf.add_run(b[1:-1]); r.italic = True
        elif b.startswith("*") and b.endswith("*"):
            scrie_inline(paragraf, b[1:-1], baza_bold=baza_bold, italic=True)
            continue
        else:
            r = paragraf.add_run(b)
        if baza_bold:
            r.bold = True
        if italic:
            r.italic = True


# --------------------------------------------------------------- tabele
def adauga_tabel(doc, linii):
    """Construieste un tabel Word dintr-un tabel markdown."""
    randuri = [[c.strip() for c in l.strip().strip("|").split("|")] for l in linii]
    aliniere = randuri[1]
    randuri = [randuri[0]] + randuri[2:]

    n_col = max(len(r) for r in randuri)
    tabel = doc.add_table(rows=len(randuri), cols=n_col)
    tabel.style = "Light Grid Accent 1"
    tabel.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, rand in enumerate(randuri):
        for j in range(n_col):
            celula = tabel.cell(i, j)
            celula.text = ""
            p = celula.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            continut = rand[j] if j < len(rand) else ""
            scrie_inline(p, continut, baza_bold=(i == 0))
            if j < len(aliniere) and aliniere[j].endswith(":"):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.size = Pt(9)
    doc.add_paragraph()
    return tabel


# ----------------------------------------------------------------- cod
def adauga_cod(doc, linii):
    for linie in linii:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(linie if linie.strip() else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    doc.add_paragraph()


# ------------------------------------------------------------ cuprins
def adauga_camp_cuprins(doc):
    """Insereaza un camp TOC pe care Word il completeaza la deschidere."""
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    inner = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Deschide in Word si apasa F9 pentru a genera cuprinsul."
    inner.append(t)
    fld.append(inner)
    p._p.append(fld)


# ------------------------------------------------------------- figuri
FIGURA_RX = re.compile(r"^>\s*\*\*(Figura \d+)\*\*\s*[—-]\s*`([^`]+)`(?:\s*\((.+?)\))?")


def adauga_figura(doc, eticheta, cale, descriere):
    if not os.path.exists(cale):
        print(f"    ! figura lipsa: {cale}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(cale, width=Cm(15.5))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    text = f"{eticheta}" + (f" — {descriere}" if descriere else "")
    r = cap.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRI
    doc.add_paragraph()


# ------------------------------------------------------- pagina de titlu
def pagina_titlu(doc, titlu, subtitlu):
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titlu); r.bold = True; r.font.size = Pt(24)

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitlu); r.font.size = Pt(13); r.italic = True
    r.font.color.rgb = GRI

    for _ in range(6):
        doc.add_paragraph()

    for text, marime, bold in [("Lucrare de practică", 14, True),
                               ("Anul III", 12, False)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.font.size = Pt(marime); r.bold = bold

    for _ in range(2):
        doc.add_paragraph()
    for eticheta in ["Student: ____________________",
                     "Coordonator: ____________________",
                     "2026"]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(eticheta); r.font.size = Pt(11); r.font.color.rgb = GRI

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ------------------------------------------------------------ antet/subsol
def numar_pagina(doc):
    subsol = doc.sections[-1].footer
    p = subsol.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for instr in ["PAGE"]:
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), instr)
        p._p.append(fld)


# ------------------------------------------------------------------ main
def converteste():
    text = open(MD, encoding="utf-8").read()
    linii = text.split("\n")

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15

    for sect in doc.sections:
        sect.top_margin = Cm(2.5); sect.bottom_margin = Cm(2.5)
        sect.left_margin = Cm(2.5); sect.right_margin = Cm(2.5)

    pagina_titlu(
        doc,
        "Predicția prețului unei locuințe folosind inteligența artificială",
        "Studiu comparativ între regresia liniară, rețelele neuronale "
        "și metodele bazate pe arbori, pe date colectate din piața "
        "imobiliară din România",
    )

    doc.add_heading("Cuprins", level=1)
    adauga_camp_cuprins(doc)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    i = 0
    n_tabele = n_figuri = n_formule = 0
    sarim_cuprins_manual = False

    while i < len(linii):
        linie = linii[i]
        gol = linie.strip()

        # titlul si subtitlul din markdown - deja pe pagina de titlu
        if gol.startswith("# ") or gol.startswith("### Studiu comparativ"):
            i += 1
            continue
        if gol.startswith("**Lucrare de practică"):
            i += 1
            continue

        # cuprinsul manual din markdown - inlocuit de campul TOC
        if gol == "## Cuprins":
            sarim_cuprins_manual = True
            i += 1
            continue
        if sarim_cuprins_manual:
            if gol.startswith("## "):
                sarim_cuprins_manual = False
            else:
                i += 1
                continue

        # separator orizontal -> pagina noua inainte de capitol
        if gol == "---":
            i += 1
            continue

        # formula pe bloc propriu
        if gol.startswith("$$"):
            bloc = []
            if gol.endswith("$$") and len(gol) > 4:
                bloc = [gol[2:-2]]
                i += 1
            else:
                i += 1
                while i < len(linii) and not linii[i].strip().endswith("$$"):
                    bloc.append(linii[i]); i += 1
                if i < len(linii):
                    bloc.append(linii[i].strip().rstrip("$$")); i += 1
            latex = " ".join(x.strip() for x in bloc if x.strip())
            cale = randeaza_formula(latex, display=True)
            if cale:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(cale)   # marime nativa: DPI-ul din PNG o face consistenta
                n_formule += 1
            continue

        # figura
        m = FIGURA_RX.match(gol)
        if m:
            adauga_figura(doc, m.group(1), m.group(2), m.group(3))
            n_figuri += 1
            i += 1
            # sarim liniile de continuare ale aceluiasi bloc citat
            while i < len(linii) and FIGURA_RX.match(linii[i].strip()):
                m2 = FIGURA_RX.match(linii[i].strip())
                adauga_figura(doc, m2.group(1), m2.group(2), m2.group(3))
                n_figuri += 1
                i += 1
            continue

        # citat simplu
        if gol.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            scrie_inline(p, gol.lstrip("> ").strip())
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = GRI
            i += 1
            continue

        # bloc de cod
        if gol.startswith("```"):
            i += 1
            bloc = []
            while i < len(linii) and not linii[i].strip().startswith("```"):
                bloc.append(linii[i]); i += 1
            i += 1
            adauga_cod(doc, bloc)
            continue

        # tabel
        if gol.startswith("|") and i + 1 < len(linii) and re.match(r"^\|[\s:|-]+\|$", linii[i+1].strip()):
            bloc = []
            while i < len(linii) and linii[i].strip().startswith("|"):
                bloc.append(linii[i]); i += 1
            adauga_tabel(doc, bloc)
            n_tabele += 1
            continue

        # titluri
        if gol.startswith("#### "):
            doc.add_heading(gol[5:].strip(), level=3); i += 1; continue
        if gol.startswith("### "):
            doc.add_heading(gol[4:].strip(), level=2); i += 1; continue
        if gol.startswith("## "):
            titlu = gol[3:].strip()
            titlu = re.sub(r"\{#.*\}", "", titlu).strip()
            doc.add_page_break()
            doc.add_heading(titlu, level=1)
            i += 1
            continue

        # lista numerotata
        m = re.match(r"^(\d+)\.\s+(.*)", gol)
        if m:
            p = doc.add_paragraph(style="List Number")
            scrie_inline(p, m.group(2)); i += 1; continue

        # lista cu puncte
        if gol.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            scrie_inline(p, gol[2:]); i += 1; continue

        # paragraf normal
        if gol:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            scrie_inline(p, gol)
        i += 1

    numar_pagina(doc)
    doc.save(OUT)
    return n_tabele, n_figuri, n_formule


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    t, f, fo = converteste()
    print(f"Generat: {OUT}")
    print(f"  tabele: {t} | figuri: {f} | formule randate: {fo}")
    print(f"  marime: {os.path.getsize(OUT)/1024:.0f} KB")
